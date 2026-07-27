"""Leitura de documentos PDF e DOCX, normalizada para texto plano UTF-8.

- PDF escaneado (sem texto) é detectado e levanta erro (OCR fora do MVP).
- DOCX com células mescladas tem o conteúdo deduplicado (atas Netcon repetem
  a mesma célula 4x devido a merge horizontal).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from loguru import logger
from pypdf import PdfReader

LIMIAR_PDF_ESCANEADO = 50


class DocumentoIlegivelError(Exception):
    """PDF escaneado ou arquivo corrompido."""


def ler_documento(caminho: Path) -> str:
    sufixo = caminho.suffix.lower()
    if sufixo == ".pdf":
        return _ler_pdf(caminho)
    if sufixo == ".docx":
        return _ler_docx(caminho)
    if sufixo == ".doc":
        raise ValueError(".doc legado não suportado — converta para .docx")
    raise ValueError(f"extensão não suportada: {sufixo}")


def _ler_pdf(caminho: Path) -> str:
    reader = PdfReader(str(caminho))
    texto = "\n".join(p.extract_text() or "" for p in reader.pages).strip()
    if len(texto) < LIMIAR_PDF_ESCANEADO:
        logger.error("PDF escaneado: {} ({} chars)", caminho.name, len(texto))
        raise DocumentoIlegivelError(f"{caminho.name} parece escaneado — OCR fora do MVP")
    return _normalizar(texto)


def _ler_docx(caminho: Path) -> str:
    doc = DocxDocument(str(caminho))
    linhas = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabela in doc.tables:
        for linha in tabela.rows:
            # dedup de células mescladas: preserva ordem, remove repetição consecutiva
            celulas: list[str] = []
            for c in linha.cells:
                txt = c.text.strip()
                if txt and (not celulas or celulas[-1] != txt):
                    celulas.append(txt)
            if celulas:
                linhas.append(" | ".join(celulas))
    texto = "\n".join(linhas).strip()
    if not texto:
        raise DocumentoIlegivelError(f"{caminho.name} vazio")
    return _normalizar(texto)


def _normalizar(texto: str) -> str:
    return (
        texto.replace("\u2022", "- ")
        .replace("\u00a0", " ")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )
